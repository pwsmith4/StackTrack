# -*- coding: utf-8 -*-
"""Build the Goodwill-facing StackTrack development overview."""

from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(r"C:\Users\Parker\OneDrive\Documents\StackTrack")
OUT = ROOT / "docs" / "StackTrack_Goodwill_Development_Overview.docx"
LOGO = ROOT / "apps" / "admin" / "src" / "assets" / "stacktrack-logo-tight.png"

NAVY = "123B5D"
BLUE = "2E74B5"
MID_BLUE = "0B6FA4"
TEAL = "008D8A"
INK = "1F2937"
MUTED = "5F7184"
LIGHT = "F2F4F7"
PALE_BLUE = "EEF7FC"
PALE_TEAL = "EFFAF8"
PALE_AMBER = "FFF7E8"
AMBER = "B7791F"
WHITE = "FFFFFF"
GRID = "D8E0E8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_table_geometry(table, widths, indent=120):
    widths = list(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_row_no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True


def shade_left_border(cell, fill=PALE_BLUE, color=MID_BLUE, width="14"):
    set_cell_shading(cell, fill)
    set_cell_border(cell, left={"val": "single", "sz": width, "space": "0", "color": color},
                    top={"val": "nil"}, bottom={"val": "nil"}, right={"val": "nil"})


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_page_field(run):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def define_numbering(doc, kind="bullet"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    lvl.extend([start, fmt, text, jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0.5)
    fmt.first_line_indent = Inches(-0.25)


def add_bullet(doc, text, num_id, style="Body Text"):
    p = doc.add_paragraph(style=style)
    apply_numbering(p, num_id)
    p.add_run(text)
    set_paragraph_spacing(p, after=5, line=1.12)
    return p


def add_numbered(doc, text, num_id, style="Body Text"):
    p = doc.add_paragraph(style=style)
    apply_numbering(p, num_id)
    p.add_run(text)
    set_paragraph_spacing(p, after=5, line=1.12)
    return p


def set_run(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_label_paragraph(doc, label, text, fill=PALE_BLUE, border=MID_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_left_border(cell, fill, border)
    p = cell.paragraphs[0]
    p.style = "Body Text"
    r = p.add_run(label.upper() + "  ")
    set_run(r, size=9, color=border, bold=True)
    r2 = p.add_run(text)
    set_run(r2, size=10.5, color=INK)
    set_paragraph_spacing(p, after=0, line=1.08)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows, widths, font_size=9.2, header_fill=LIGHT, zebra=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    set_row_no_split(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        clear_paragraph(p)
        r = p.add_run(value)
        set_run(r, size=8.5, color=NAVY, bold=True)
        set_paragraph_spacing(p, after=0, line=1.0, keep=True)
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        set_row_no_split(row)
        for cidx, value in enumerate(row_data):
            cell = row.cells[cidx]
            if zebra and ridx % 2 == 1:
                set_cell_shading(cell, "FAFCFE")
            p = cell.paragraphs[0]
            clear_paragraph(p)
            if isinstance(value, tuple):
                text, color, bold = value
                r = p.add_run(str(text))
                set_run(r, size=font_size, color=color, bold=bold)
            else:
                r = p.add_run(str(value))
                set_run(r, size=font_size, color=INK)
            set_paragraph_spacing(p, after=0, line=1.03)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                            top={"val": "single", "sz": "4", "space": "0", "color": GRID},
                            bottom={"val": "single", "sz": "4", "space": "0", "color": GRID},
                            left={"val": "single", "sz": "4", "space": "0", "color": GRID},
                            right={"val": "single", "sz": "4", "space": "0", "color": GRID})
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_paragraph_spacing(p, before={1:16, 2:12, 3:8}.get(level, 8), after={1:8, 2:6, 3:4}.get(level, 4), line=1.05, keep=True)
    return p


def add_body(doc, text, style="Body Text", after=6):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    set_paragraph_spacing(p, after=after, line=1.1)
    return p


def add_small(doc, text):
    p = doc.add_paragraph(style="Small")
    p.add_run(text)
    set_paragraph_spacing(p, after=4, line=1.0)
    return p


def add_section_rule(doc, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, NAVY, 8, 4)):
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing = 1.05
        s.paragraph_format.keep_with_next = True
    for name, base, size, color in (("Body Text", "Normal", 10.8, INK), ("Small", "Normal", 8.7, MUTED), ("Kicker", "Normal", 9, BLUE), ("Subtitle", "Normal", 13, MUTED), ("Callout", "Normal", 10.2, NAVY)):
        if name in styles:
            s = styles[name]
        else:
            s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles[base]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.1
    styles["Kicker"].font.bold = True
    styles["Kicker"].paragraph_format.space_after = Pt(4)
    styles["Subtitle"].paragraph_format.space_after = Pt(10)
    styles["Callout"].font.bold = True


def configure_section(section):
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("STACKTRACK  /  GOODWILL DEVELOPMENT OVERVIEW")
    set_run(r, size=8, color=MUTED, bold=True)
    set_paragraph_spacing(p, after=0, line=1.0)
    footer = section.footer
    p = footer.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Goodwill pilot brief  |  July 31, 2026  |  Page ")
    set_run(r, size=8, color=MUTED)
    r2 = p.add_run()
    set_run(r2, size=8, color=MUTED)
    add_page_field(r2)
    set_paragraph_spacing(p, after=0, line=1.0)


def add_cover(doc, bullet_id):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run()
        r.add_picture(str(LOGO), width=Inches(2.15))
        set_paragraph_spacing(p, after=16, line=1.0)
    p = doc.add_paragraph(style="Kicker")
    p.add_run("GOODWILL  |  REUSABLE CONTAINER OPERATIONS")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("StackTrack\nDevelopment Overview")
    set_run(r, size=29, color=NAVY, bold=True)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Current pilot capabilities, accuracy model, operating workflows, and the path to production.")
    add_section_rule(doc, BLUE)
    meta_rows = [
        ("Prepared for", "Goodwill operations and technology stakeholders"),
        ("Prepared by", "Parker Smith / StackTrack development"),
        ("Document date", "July 31, 2026"),
        ("Current status", ("Working test pilot", TEAL, True)),
        ("Data boundary", ("Synthetic test data only; no live Goodwill data", AMBER, True)),
    ]
    add_table(doc, ["Document field", "Current value"], meta_rows, [2200, 7160], font_size=10, header_fill=PALE_BLUE, zebra=False)
    add_label_paragraph(doc, "Read this first", "The first half describes features that are implemented and testable now. The second half identifies production prerequisites, decisions Goodwill must confirm, and a staged rollout plan. A feature marked planned is not represented as production-ready.", fill=PALE_AMBER, border=AMBER)
    add_heading(doc, "Purpose of this brief", 2)
    add_body(doc, "StackTrack is an accuracy-first system for tracking reusable bins, carts, and gaylords as they are filled, sent out, received, and emptied. It replaces unreliable paper handoffs with a durable event history, a current-state view, an offline-capable mobile workflow, and administrator tools for investigation and governed correction.")
    add_body(doc, "This overview is designed to help Goodwill evaluate the working foundation, see how the pieces fit together, and decide what must be confirmed before a real pilot. It is a development overview, not a final statement of Goodwill policy, a service-level agreement, or an authorization to enter live operational data.")
    add_heading(doc, "How to read the status language", 3)
    add_bullet(doc, "Implemented / pilot-ready means the behavior exists in the current test branch and is covered by the local or Azure synthetic test workflow.", bullet_id)
    add_bullet(doc, "Planned hardening means the product direction is known but must be connected to Goodwill identity, hardware, networking, policies, or production integrations.", bullet_id)
    add_bullet(doc, "Decision required means Goodwill operating owners need to define the rule before it can be safely treated as official behavior.", bullet_id)


def add_contents(doc):
    add_heading(doc, "Contents", 1)
    contents = [
        ("1", "Executive summary"), ("2", "Product goals and design principles"), ("3", "System at a glance"),
        ("4", "Mobile scanner experience"), ("5", "Administrator console"), ("6", "Accuracy and data model"),
        ("7", "Governance, roles, and security"), ("8", "Reporting, inventory, and planning"),
        ("9", "Representative operational scenarios"), ("10", "Test environment and deployment model"),
        ("11", "Current implementation versus production readiness"), ("12", "Recommended roadmap"),
        ("13", "Goodwill discovery questions and decisions"), ("14", "Risks, controls, and acceptance criteria"),
        ("15", "Glossary"),
    ]
    add_table(doc, ["Section", "Topic"], contents, [1100, 8260], font_size=10, zebra=True)
    add_label_paragraph(doc, "One important business decision", "A location employee does not know the truck's destination. The mobile app therefore records the departure origin only. A receiving location is recorded when that location scans the arrival; StackTrack must not present a guessed destination as an official plan.", fill=PALE_TEAL, border=TEAL)


def add_exec_summary(doc, bullet_id):
    add_heading(doc, "1. Executive summary", 1)
    add_body(doc, "The current StackTrack foundation is a working, end-to-end test system: a shared-device mobile app records container observations; a Fastify API validates and stores events; PostgreSQL provides the durable system of record; and a React administrator website turns the evidence into operational views, review queues, reports, and controlled administrative actions.")
    add_body(doc, "The design is intentionally conservative about accuracy. StackTrack preserves observations rather than overwriting them, separates physical evidence from the current projection, keeps device and server timing evidence, queues work while a device is offline, and routes ambiguous events to review. This makes the system explainable when a scan is late, duplicated, out of order, or inconsistent with the previous state.")
    add_table(doc, ["What the pilot does now", "Why it matters to Goodwill"], [
        ("Records fill, departure, arrival, and empty events from a shared scanner.", "Creates a common operating language that can replace paper handoffs."),
        ("Works through offline capture and later replay.", "A temporary loss of connectivity does not force employees back to paper."),
        ("Shows the latest usable state while retaining the full event history.", "Admins can act on the current picture without losing the evidence behind it."),
        ("Explains flags and routes in human language.", "A manager can understand what happened without reading database codes."),
        ("Provides review, correction, device, location, user, audit, and export controls.", "Corporate operations can govern exceptions without silently rewriting history."),
    ], [4200, 5160], font_size=9.7)
    add_heading(doc, "What is deliberately not claimed yet", 2)
    for text in [
        "The test deployment is not connected to live Goodwill data, live production systems, or production identity. Its records are synthetic and safe for demonstrations only.",
        "The mobile workflow currently uses typed or keyboard-wedge label entry. Camera/QR behavior and physical Unitech scanner integration remain pilot work.",
        "Warehouse outlook and service planning are useful scenario tools today, but their forecasts require real historical data and Goodwill-approved operating targets before they should drive dispatch decisions.",
        "Destination is not known at departure. The system records “departed from” and closes the handoff only when a receiving location records the arrival.",
    ]:
        add_bullet(doc, text, bullet_id)


def add_principles(doc, bullet_id):
    add_heading(doc, "2. Product goals and design principles", 1)
    add_table(doc, ["Principle", "How StackTrack applies it", "Operational benefit"], [
        ("Accuracy first", "Never silently erase a scan; preserve evidence and show why a state is trusted or needs review.", "A discrepancy can be investigated without guessing what the system changed."),
        ("Offline first", "Persist an observation locally with a stable event ID and sequence, then replay it safely when connected.", "Connectivity loss becomes a visible queue, not an invisible data loss."),
        ("Human-readable by default", "Translate technical flags, event names, route state, and review reasons into plain-language labels.", "Store and corporate users can understand the action without code knowledge."),
        ("Least privilege", "Scope location managers to assigned locations; reserve cross-location governance and user management for corporate roles.", "Local teams can keep work moving without gaining unrelated network access."),
        ("Auditability", "Administrative controls record actor, target, before/after context, reason, and decision time.", "Goodwill can answer who changed an operational control and why."),
        ("Replaceable adapters", "Keep scanner, identity, cloud, and production-system connections behind clear adapters.", "The pilot can prove business rules before hardware and enterprise integrations are finalized."),
    ], [1900, 4400, 3060], font_size=9.2)
    add_heading(doc, "Accuracy is a behavior, not a badge", 2)
    add_body(doc, "A green or clean state means the current projection has no active evidence-quality issue under the configured rules. It does not prove that a container is physically present beyond all doubt. A warning means the event remains usable but deserves attention; a review state means the system is deliberately withholding an official projection until an authorized person decides what the evidence supports.")
    add_label_paragraph(doc, "Plain-language rule", "The console should tell an administrator what happened, what is known, what is not known, and what action is available. Technical identifiers remain available in Details for traceability, but they are not the primary story on screen.", fill=PALE_BLUE, border=BLUE)


def add_architecture(doc, bullet_id):
    add_heading(doc, "3. System at a glance", 1)
    add_body(doc, "The system has four operational surfaces and one durable evidence path. The website and mobile app are clients; the API owns validation and authorization; PostgreSQL stores the event ledger and projections. Reports read the same evidence rather than maintaining a competing set of counts.")
    add_table(doc, ["Layer", "Current implementation", "Production direction"], [
        ("Shared scanner app", "Expo/React Native workflow with label entry, four actions, messages, offline queue, replay, refresh, device identity, and version telemetry.", "Managed Android build, camera/QR or Unitech adapter, MDM/Intune provisioning, employee session policy."),
        ("API", "Fastify/TypeScript routes for events, state, reference data, devices, review, corrections, admin sessions, reports, and exports.", "Azure Container Apps or equivalent with private configuration, monitoring, production identity, and controlled CORS."),
        ("System of record", "PostgreSQL schema with tenant isolation, row-level security foundation, idempotency constraints, and append-only enforcement.", "Azure Database for PostgreSQL with backup, retention, alerting, private networking, and tested recovery."),
        ("Admin website", "React/Vite console published separately from the API, with overview, inventory, locations, activity, audit, devices, reports, service planning, outlook, and settings.", "Goodwill Entra sign-in, role mapping, accessibility review, corporate branding decisions, and production deployment controls."),
        ("Analytics boundary", "CSV exports and prototype reporting/forecast views; Microsoft analytics connection is a planned boundary.", "Governed export or Fabric/ADLS/Power BI integration after data definitions and security review."),
    ], [1900, 4300, 3160], font_size=9.2)
    add_heading(doc, "Evidence path", 2)
    add_table(doc, ["Step", "What happens"], [
        ("1. Observe", "A shared scanner identifies a container, action, location, device, event time, sequence, and optional goods/message details."),
        ("2. Queue if needed", "If offline, the mobile app stores the complete event locally with its UUID and status."),
        ("3. Validate", "The API checks the tenant, device assignment, idempotency, event shape, reference version, and ordering evidence."),
        ("4. Append", "The accepted observation is added to the immutable ledger. A duplicate replay is handled idempotently."),
        ("5. Project", "The system derives the latest usable container state and open handoff from the evidence; ambiguity creates a review case."),
        ("6. Explain", "Admin pages, Activity, Audit, review queues, reports, and exports present the same evidence for different questions."),
    ], [1600, 7760], font_size=9.6)
    add_label_paragraph(doc, "Boundary to remember", "GitHub Pages hosts the static admin client only. It does not run the API or PostgreSQL. The API is deployed separately, and the browser must be authenticated and allowed by the API before operational data is returned.", fill=PALE_AMBER, border=AMBER)


def add_mobile(doc, bullet_id):
    add_heading(doc, "4. Mobile scanner experience", 1)
    add_body(doc, "The mobile application is designed for shared scanners. The physical device is locked to an assigned location and carries a short five-digit operator-facing scanner ID. A scan identifies the container and action; the application does not ask a location employee to choose a truck destination they cannot know.")
    add_heading(doc, "Current scan workflow", 2)
    add_table(doc, ["Employee action", "What the app records", "What the employee sees"], [
        ("Mark full", "Container, current assigned location, fill event, goods category and classification, device identity, event time, and optional note.", "The bin/cart/gaylord is marked full and a load code is generated for the downstream production handoff."),
        ("Send out", "Container, departure origin, device identity, event time, and optional note. No destination is required or implied.", "The container is shown as departed from the current location; receiving site not yet confirmed."),
        ("Receive", "Container and receiving location at the device's assigned site, with device/time evidence.", "The container is received here. The system can explain the source from the open departure history when the route is unambiguous."),
        ("Mark empty", "Container, current location, device identity, event time, and optional note.", "The container is available/empty at the location for the next operating decision."),
        ("Message/note", "A human note linked to the observation and container.", "The note appears in the admin container, Activity, and relevant report/detail views."),
    ], [1700, 4500, 3160], font_size=9.1)
    add_heading(doc, "Offline behavior", 2)
    for text in [
        "The app persists the observation locally before showing it as pending. It does not rely on a network response to protect the scan from immediate loss.",
        "Each event has a stable UUID and a monotonic device sequence. Replaying the same event is safe; replaying a different payload under the same UUID is an integrity failure rather than an overwrite.",
        "When connectivity returns, the app retries automatically and also provides an explicit refresh/retry action. The employee can distinguish pending, synchronized, and review-needed results.",
        "A disabled scanner cannot create or upload operational events. The admin console records the device control and the mobile app makes the operational consequence visible.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_heading(doc, "Version and device identity", 2)
    add_body(doc, "The build has a StackTrack app version and reports it with device telemetry. The admin Devices page shows the five-digit scanner ID, label, assigned location, availability, installed version (when reported), observation count, pending/offline state, and last report. Releases can update the mobile version automatically in GitHub artifacts, but installing or managing the new build on physical devices remains a deployment/MDM responsibility.")
    add_label_paragraph(doc, "Pilot hardware boundary", "Typed label entry and an Expo/Android emulator prove the workflow and data rules. Goodwill still needs to confirm the exact Unitech model, scanning interface, camera/QR expectations, Android version, charging practice, MDM, and replacement process before hardware acceptance.", fill=PALE_BLUE, border=BLUE)


def add_admin(doc, bullet_id):
    add_heading(doc, "5. Administrator console", 1)
    add_body(doc, "The admin console is the corporate operational workspace. It is organized around the questions an administrator needs to answer: what is happening now, where is a container, what needs attention, who changed an operational control, and what should transportation or site leadership do next?")
    page_rows = [
        ("Overview", "Network pulse, current container footprint, active handoffs, review attention, location operations, inventory summary, warehouse snapshot, and links into detailed work.", "Daily triage and network awareness."),
        ("Company-wide inventory", "Location-by-container-type matrix, goods mix, current/on-hand counts, filters, and export. Current data is state/projection based; future history snapshots support trend comparison.", "Company-wide inventory and reconciliation."),
        ("Daily service plan", "Location and goods/container filters, pickup and empty-crate delivery priorities, target setup, service date, priority queue, and export.", "Transportation planning from operating targets."),
        ("Warehouse outlook", "Forecast horizon, warehouse/store coverage, target setup, holiday and planning adjustments, calculation explanation, history/trend view, and export.", "Scenario planning, not an official scan correction."),
        ("Containers", "Detailed container search and filters by state, type, locations involved, movement, health, message status, time, and sort. Row/details view includes current projection, route context, history, and messages.", "Locate and investigate one or many assets."),
        ("Load codes", "Validated load-code list with location, container, goods, time, state, filters, details, and export of the current result set.", "Reconcile production-facing handoffs."),
        ("Locations", "Recommended network view with location-type icons, current counts, open handoffs, scans, review count, filters, export, and deep links to a focused location workspace.", "Network operations and local drill-down."),
        ("Location workspace", "A selected store, Donation Xpress, or warehouse shows local scanners, local activity, review items, current containers, departures, and local health. The page does not claim to know inbound destinations before arrival.", "Scoped local operations."),
        ("Needs review / Corrections", "Review cases, evidence, reasons, approval/rejection/reopen decisions, and append-only correction history.", "Resolve ambiguity without erasing evidence."),
        ("Activity", "Physical scanner observations and movement with human event narratives, exact action/location/scanner/time filters, related-event grouping, and links to container details.", "What physically happened?"),
        ("Audit trail", "Administrator sign-ins, device controls, user/location changes, approvals, corrections, reasons, before/after context, detailed filters, pagination, and export.", "Who changed the system, when, and why?"),
        ("Devices", "Search and filter scanners by location, enabled state, version, freshness, and other health signals. Rename, move, enable/disable, view assignment history, and inspect details.", "Manage shared scanner fleet."),
        ("Reports & data", "Read-only operational reports, data-health signals, filterable evidence, and CSV exports for movement, load-code, exceptions, corrections, scanner coverage, location throughput, transit aging, latency, and governance actions.", "Investigation and reproducible handoff."),
        ("Settings", "Access/user management, roles, location administration, scanner policies, correction policy, integration boundaries, and account security.", "Low-frequency governance and configuration."),
    ]
    add_table(doc, ["Console area", "Current capability", "Primary question"], page_rows, [1800, 5100, 2460], font_size=8.7)
    add_heading(doc, "Details are for decisions, not database inspection", 2)
    add_body(doc, "Detail drawers and pages intentionally lead with plain-language status, location, route context, timestamps, messages, available actions, and next steps. Technical IDs remain available through copy/details controls for support and audit traceability. Raw JSON and internal event names are not presented as the primary user experience.")
    add_heading(doc, "Overview information that earns its space", 2)
    add_body(doc, "The overview is designed to remain a useful daily landing page rather than becoming a second reports page. Its core sections are: network counts and attention; a location network view; active handoffs using origin-only language; recent physical activity; company-wide inventory summary; warehouse/store planning snapshot; and direct links to review, service, and outlook work. Deep histories, complex filters, and configuration live on their own pages.")


def add_accuracy(doc, bullet_id):
    add_heading(doc, "6. Accuracy and data model", 1)
    add_body(doc, "StackTrack separates three things that are often conflated in paper processes: what a scanner observed, what the system currently projects, and what an administrator later decided. That separation is the foundation for explainable accuracy.")
    add_heading(doc, "Core event language", 2)
    add_table(doc, ["Employee-facing action", "Internal event concept", "Projection effect"], [
        ("Marked full", "load_assigned", "Container is loaded/full at the observed location and a load-code record can be generated."),
        ("Departed", "batch_out", "Container leaves the origin and enters an open handoff; destination remains unknown until a receipt."),
        ("Arrived", "batch_in", "Receiving location becomes the current location; source can be narrated from an unambiguous open handoff."),
        ("Marked empty", "emptied", "Container becomes empty/available at the observed location."),
    ], [2300, 2100, 4960], font_size=9.4)
    add_heading(doc, "The three timestamps", 2)
    add_table(doc, ["Timestamp", "Meaning", "Why it is kept"], [
        ("Device observed", "When the scanner says the physical action occurred.", "Preserves the device's original evidence, even when it was offline or its clock was imperfect."),
        ("Effective", "The time used for projection/reporting after the measured device/server offset is considered.", "Makes ordering more reliable without throwing away raw evidence."),
        ("Server received", "When the API accepted the upload.", "Measures offline/sync latency and distinguishes late upload from late physical action."),
    ], [1800, 4300, 3260], font_size=9.3)
    add_heading(doc, "Conflict and replay rules", 2)
    for text in [
        "Same event UUID and same payload: treat as an idempotent replay; do not create a second observation.",
        "Same event UUID and different payload: reject as an integrity failure and retain enough context for investigation.",
        "Duplicate or out-of-order device sequences: preserve the observations, flag the evidence, and avoid presenting an overconfident state.",
        "Clock offset: retain the raw device timestamp; use the measured offset for effective ordering; warn or review when thresholds are exceeded.",
        "Contradictory observations: retain both; create a review case; do not silently replace history with the latest upload.",
        "Approved corrections: add a newer, attributed administrative decision; never update or delete the original observation.",
        "Later physical evidence can supersede an administrative projection when the newer event is accepted under the normal rules.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_heading(doc, "Unknown destination is a first-class state", 2)
    add_body(doc, "Because the person at the departure site does not know where the truck will go, the system must use “departed from [origin]” and “receiving site not yet confirmed.” It should not show a guessed destination, a planned route, or a directional arrow as if it were official. Once an arrival scan occurs, the system can show “arrived at [receiving location] from [source]” and reconstruct multi-hop movement such as Donation Xpress → Warehouse A → Warehouse B → Store from the observed sequence.")
    add_label_paragraph(doc, "What this protects", "An item in transit is not missing simply because its destination is unknown. It is an open handoff with a confirmed origin and an unconfirmed receiving location. That distinction prevents false certainty while still giving transportation a useful aging and follow-up signal.", fill=PALE_TEAL, border=TEAL)


def add_governance(doc, bullet_id, num_id):
    add_heading(doc, "7. Governance, roles, and security", 1)
    add_body(doc, "The recommended operating model has corporate governance above location operations. Goodwill should own the Organization Owner role; no hidden developer super-admin should exist in the production system. A second owner is recommended for continuity and dual control over material decisions.")
    add_table(doc, ["Role", "Scope", "Can do", "Cannot do"], [
        ("Organization Owner", "Goodwill corporate / network-wide", "Manage users and location scope; configure policy; approve/reject/reopen corrections; manage devices and locations; view all reports and audit.", "Erase immutable evidence; bypass required dual control; use a hidden vendor backdoor."),
        ("Operations Administrator", "Network-wide operational administration", "Manage scanners; investigate containers and reviews; request corrections; run reports and service planning; maintain normal operational settings.", "Add/demote administrators; approve their own material correction; access outside assigned scope if Goodwill chooses scoped admin."),
        ("Location Manager", "Assigned store, Donation Xpress, or warehouse", "View local containers/activity/reviews; manage local scanner availability/name; submit reasoned corrections or requests.", "Approve own correction; change other locations; add admins; erase observations; change corporate policy."),
        ("Read-only Reviewer", "All or assigned locations", "View agreed operational/reporting evidence and export if allowed.", "Change devices, users, locations, corrections, or policy."),
        ("StackTrack Support", "Time-limited, explicitly approved", "Troubleshoot under an auditable grant; default read-only.", "Permanent access, hidden ownership, or unreviewed production changes."),
    ], [1700, 2200, 3350, 2110], font_size=8.4)
    add_heading(doc, "Correction and approval path", 2)
    for text in [
        "A local or operations user identifies a specific evidence problem and provides a clear reason, target, affected location, and supporting observation(s).",
        "The request is append-only and appears in a corporate review queue; the original scan stays intact.",
        "An authorized owner approves, rejects, or reopens the request. A material change requires a different owner from the requester.",
        "The decision, actor, reason, before/after context, and decision time are written to the audit trail.",
        "A later accepted physical scan can supersede the correction under the normal projection rules.",
    ]:
        add_numbered(doc, text, num_id)
    add_heading(doc, "Pilot authentication versus production authentication", 2)
    add_table(doc, ["Pilot test", "Production requirement"], [
        ("Server-issued administrator session with a bootstrap owner; password hash only; public site hides tenant data until a verified session.", "Goodwill Microsoft Entra ID/OIDC sign-in with MFA/conditional access, mapped to a server-side tenant user and role."),
        ("Fixed development device/tenant identity for synthetic testing.", "Provisioned device identity, per-device credentials or certificates, MDM enrollment, revocation, and replacement workflow."),
        ("CORS and rate limits are implemented for the test surface.", "Goodwill-approved domain allowlist, secret manager, monitoring, alerting, penetration review, and incident response."),
        ("Test password reset/session controls exist for the pilot bridge.", "Password lifecycle is owned by Entra; StackTrack never displays passwords and records only administrative control events."),
    ], [4680, 4680], font_size=9.0)
    add_label_paragraph(doc, "Recommended ownership", "You should be the initial Organization Owner for development and pilot administration only if Goodwill explicitly authorizes it. At sale/rollout, Goodwill's Chief of IT or designated corporate administrator should be an Organization Owner, with at least one additional Goodwill owner for continuity and dual control.", fill=PALE_AMBER, border=AMBER)


def add_reporting(doc, bullet_id):
    add_heading(doc, "8. Reporting, inventory, and planning", 1)
    add_body(doc, "Reporting is an investigation and planning layer. It reads immutable observations, current projections, scanner telemetry, correction decisions, and audit records. It does not become a second source of truth, and a CSV export is read-only evidence rather than a mechanism to rewrite the system.")
    add_heading(doc, "Reports and data", 2)
    add_table(doc, ["Report", "Use it for", "Do not interpret it as"], [
        ("Movement ledger", "Reconstruct observations, locations, scanners, event times, receipt times, and out-of-order evidence.", "Proof that every physical movement was correct."),
        ("Load-code handoff", "Reconcile full containers and production-facing load codes for a daily handoff.", "Proof that the destination received the container."),
        ("Data-quality exceptions", "Prioritize flags, unresolved conflicts, clock/sequence issues, and missing evidence.", "A list of assets that are necessarily lost."),
        ("Correction register", "Review requests, decisions, reasons, before/after context, and re-opened items.", "A replacement for the immutable observation ledger."),
        ("Scanner coverage", "See enabled/disabled devices, location assignment, version, freshness, and registration coverage.", "Proof that a quiet location had no physical activity."),
        ("Location throughput", "Compare observations and distinct containers by location and time window.", "A simple performance score; volume alone is not quality."),
        ("Transit aging", "Find open handoffs with a confirmed departure origin, elapsed time, and no receiving scan.", "A known destination or automatic service-level breach."),
        ("Scan latency", "Separate offline upload delay from device/service behavior.", "Automatic evidence that the physical scan was inaccurate."),
        ("Governance actions", "Review administrator controls, approvals, and configuration changes.", "The physical activity feed; use Activity for that."),
    ], [1850, 4300, 3210], font_size=8.6)
    add_heading(doc, "Company-wide inventory", 2)
    add_body(doc, "The inventory workspace is intended to reproduce the useful shape of Goodwill's current spreadsheet/report: locations down the side, container types and goods categories across the top, totals, filters, and export. The current pilot can show current counts by location/type and export CSV that opens in Excel. Native XLSX formatting, weekly historical snapshots, and validated trend definitions should be added after Goodwill confirms the official inventory vocabulary and reporting cadence.")
    add_heading(doc, "Daily service plan", 2)
    add_body(doc, "The service plan turns configured minimum/maximum or target ranges into a review queue: which sites appear to need pickup of full crates, which need delivery of empty crates, which are critically short, and which rows need data verification. It is a recommendation for transportation review, not an automatic dispatch order. The service date, location, goods category, container type, status, and export are intended to make the queue reproducible.")
    add_heading(doc, "Warehouse outlook", 2)
    add_body(doc, "Warehouse outlook is a separate planning workspace rather than a crowded overview widget. It supports an adjustable receipt horizon (for example, next 7/14/30 days), warehouse and goods filters, current target setup, store coverage, history/trend context, holiday/seasonal planning inputs, and a calculation explanation. The pilot uses synthetic history and scenario inputs. Goodwill must validate the forecast approach against real historical receipts, holidays, capacity constraints, transportation schedules, and stockout outcomes before using it as an operational commitment.")
    add_table(doc, ["Planning input", "Current direction", "Goodwill decision needed"], [
        ("Minimum / maximum", "Target setup can be adjusted by location and container/goods mix.", "What is the true minimum safe on-hand amount and maximum useful buffer for each store and type?"),
        ("Expected receipts", "Forecast horizon is adjustable and clearly labeled as a time window.", "What counts as expected: open departures, historical receipt pattern, scheduled route, or a combination?"),
        ("Holiday/seasonality", "Holiday adjustments can be represented as scenario multipliers/inputs.", "Which holidays/events matter, what lead time applies, and who owns the adjustment?"),
        ("Planning gap", "The view can explain how expected, recommended on-hand, and gap are derived.", "Which operating formula should be authoritative and how should exceptions be handled?"),
        ("Transport priority", "Service rows are ranked by shortage/urgency signals.", "What makes a row critical, and who can override the recommendation?"),
    ], [2200, 3600, 3560], font_size=8.8)
    add_heading(doc, "Data health in plain language", 2)
    add_table(doc, ["Signal", "Specific use case", "Suggested response"], [
        ("Observation integrity", "Timing, sequence, or device-order evidence is questionable.", "Review the scan details and decide whether the event needs a governed correction."),
        ("Projection decisions", "The system cannot safely choose one official state from conflicting evidence.", "Open the review queue; do not assume the newest upload is correct."),
        ("Scanner freshness", "A device has not reported recently.", "Check connectivity, power, assignment, app status, and whether the site actually had activity."),
        ("Registration coverage", "A registered container has not produced accepted evidence.", "Confirm whether it exists, is labeled, is deployed, or should be retired."),
        ("Upload latency", "An event was observed earlier but arrived late.", "Investigate offline duration/network/device behavior; do not call the scan wrong solely because it was late."),
    ], [2100, 4300, 2960], font_size=8.9)


def add_scenarios(doc, bullet_id):
    add_heading(doc, "9. Representative operational scenarios", 1)
    add_heading(doc, "A. Normal fill and departure", 2)
    add_table(doc, ["Moment", "Mobile action", "Admin interpretation"], [
        ("At store", "Employee scans B1001, chooses Mark full, selects goods/classification, optionally enters a message, and confirms.", "B1001 is full at the store; a load code is available for the production-facing handoff."),
        ("Truck leaves", "Employee scans Send out; no destination is selected.", "B1001 departed from the store. Receiving location is not confirmed."),
        ("At receiving site", "Receiving employee scans Receive on a device assigned to that site.", "B1001 arrived at the receiving site; the route can be narrated from the departure origin."),
    ], [1800, 4300, 3260], font_size=9.0)
    add_heading(doc, "B. Multi-hop movement", 2)
    add_body(doc, "A container may go Donation Xpress → Warehouse A → Warehouse B → Store. Each departure records the location it left, and each arrival records the location that received it. The console should present the confirmed segments in chronological order, not invent a single simple route. Until an arrival is scanned, the object is open in transit with origin known and destination unknown.")
    add_heading(doc, "C. Offline capture", 2)
    for text in [
        "The device loses connectivity after a worker marks a container full. The app stores the scan as pending and keeps it visible locally.",
        "The worker can continue using the workflow subject to the device's local queue/disabled rules; the admin cannot see the event until it is uploaded.",
        "Connectivity returns. The app retries/syncs, the API accepts the event once, and the admin Activity/Containers/Load Codes pages show it after refresh.",
        "If the replay exposes a sequence or timing problem, the event remains evidence and is routed to review rather than silently discarded.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_heading(doc, "D. Conflicting or missing evidence", 2)
    add_body(doc, "If two devices report incompatible events or an arrival appears without a clear departure, StackTrack keeps the records and presents a review case. An authorized user can request or approve a correction with a reason. The correction is a new administrative record; it does not change the immutable scan history.")
    add_heading(doc, "E. Scanner maintenance", 2)
    add_body(doc, "An administrator can disable a broken scanner, rename it, move it to another location, inspect its assignment history, and review its app version/freshness. Low-risk controls can take effect immediately with an audit record; cross-location or policy-sensitive changes should be governed by role and Goodwill approval rules.")


def add_deployment(doc, bullet_id):
    add_heading(doc, "10. Test environment and deployment model", 1)
    add_body(doc, "The current environment is intentionally close enough to production to test the boundaries while remaining safe. It uses synthetic data, a separately deployed API, PostgreSQL, and a static admin site. It is not a live Goodwill operating environment.")
    add_table(doc, ["Surface", "Where it runs in testing", "What it proves"], [
        ("Admin website", "GitHub Pages static build from the test branch.", "Navigation, filters, details, role-aware UI, exports, and browser behavior."),
        ("API", "Azure Container App test deployment from a public GHCR image for synthetic testing.", "Real HTTP/auth/session behavior, device controls, reviews, reports, and database connectivity."),
        ("Database", "Azure Database for PostgreSQL Flexible Server test instance (or local PostgreSQL for offline development).", "Schema, projections, tenant isolation foundation, persistence, and database constraints."),
        ("Mobile", "Expo preview, Android emulator, or local debug build pointed at the test API.", "Shared-device workflow, offline queue, version/device telemetry, and replay."),
    ], [1700, 4200, 3460], font_size=9.1)
    add_heading(doc, "Local development from a clean restart", 2)
    add_body(doc, "The repository includes a Windows `start-local.cmd` shortcut for the local API, admin, and mobile preview. Local PostgreSQL can be started and seeded separately. The local guide also documents Android emulator port forwarding and the cloud-mobile path, which uses the Azure test API rather than a local API.")
    add_heading(doc, "What is currently safe and unsafe", 2)
    add_table(doc, ["Safe for the test environment", "Not safe to treat as production"], [
        ("Synthetic containers, locations, scanners, observations, review cases, exports, and scenario planning inputs.", "Real Goodwill labels, employee identities, customer/personally identifiable information, or live production counts."),
        ("Testing offline capture, replay, conflict handling, filters, reports, and admin workflows.", "Using the pilot password bridge as Goodwill's long-term identity system."),
        ("Testing API/database deployment, CORS, rate limiting, and CI/CD behavior.", "Assuming a public static site plus a test container is a complete security, backup, or availability design."),
    ], [4680, 4680], font_size=9.1)
    add_label_paragraph(doc, "Data migration", "No data migration is currently implemented or assumed. Goodwill will need to decide whether the pilot starts with a clean opening inventory, a label registration import, an opening balance, or a controlled migration from paper/spreadsheets.", fill=PALE_AMBER, border=AMBER)


def add_readiness(doc, bullet_id):
    add_heading(doc, "11. Current implementation versus production readiness", 1)
    add_body(doc, "The following matrix is intended to keep a Goodwill review honest. “Working” means it can be exercised in the current test environment. “Production gate” means additional Goodwill decisions or engineering hardening are still required.")
    rows = [
        ("Event contracts and projection", "Working", "Run real-world scenario tests; confirm official event semantics and thresholds."),
        ("Offline queue and idempotent replay", "Working", "Test worst-case outage length, queue limits, device replacement, and duplicate behavior."),
        ("Unknown-destination movement model", "Working direction", "Confirm departure/arrival operating procedure and multi-hop reporting language."),
        ("Admin operations and reports", "Working pilot", "Accessibility, role-scope review, export definitions, and performance at Goodwill scale."),
        ("Device administration", "Working pilot", "Choose hardware identity/provisioning, MDM, credentials/certificates, and replacement flow."),
        ("Owner/admin/location roles", "Working pilot", "Map to Goodwill Entra groups, formalize approval thresholds, and validate scoped access."),
        ("Pilot password session", "Test only", "Replace with Entra ID/MFA and secret management before real data."),
        ("PostgreSQL tenant/RLS foundation", "Implemented foundation", "Security review, private networking, backups, restore drills, monitoring, and scale test."),
        ("Load-code integration", "Local/test behavior", "Confirm production system interface, ownership, validation, retries, and reconciliation."),
        ("Warehouse forecasts", "Scenario prototype", "Load real history; validate formulas, holidays, capacity, and planner trust."),
        ("Camera/Unitech integration", "Not complete", "Hardware proof-of-concept, Android build, scanning reliability, and MDM rollout."),
        ("Analytics/Fabric/data lake", "Boundary / planned", "Define governed data products, retention, access, cost, and refresh expectations."),
        ("24/7 operational support", "Not included in pilot", "Define support ownership, hours, response targets, incident path, and commercial model."),
    ]
    add_table(doc, ["Capability", "Current state", "Next gate"], rows, [2500, 1900, 4960], font_size=8.8)
    add_heading(doc, "Automated verification already in place", 2)
    add_body(doc, "As of July 31, 2026, the current test branch passes 48 automated tests across 11 test files, passes TypeScript checks for the admin, API, mobile, domain, and offline-queue workspaces, and produces a successful Vite admin build. This is useful evidence of a stable development baseline; it is not a substitute for Goodwill field validation, security review, performance testing, or hardware acceptance.")


def add_roadmap(doc, bullet_id, num_id):
    add_heading(doc, "12. Recommended roadmap", 1)
    add_body(doc, "The safest path is to finish the rules and pilot workflow before adding every integration. Each phase should end with an observable acceptance decision, not just a code merge.")
    phases = [
        ("Phase 0 — Goodwill discovery", "Confirm container taxonomy, label ownership, event meanings, unknown-destination process, location ownership, correction policy, role model, offline expectations, hardware, and pilot success measures.", "Approved operating rules and a signed pilot scope."),
        ("Phase 1 — Production foundation", "Entra ID/MFA, Goodwill tenant and roles, secret manager, private networking, backups/restore, observability, CORS/domain controls, rate limits, audit retention, and deployment gates.", "A security-reviewed environment with no pilot password bridge."),
        ("Phase 2 — Hardware and two-location pilot", "Integrate the chosen scanners, deploy managed builds, run store/warehouse workflows, exercise offline/replay, train users, and measure scan accuracy and adoption.", "Goodwill can operate the workflow for an agreed test window without reverting to paper."),
        ("Phase 3 — Operational integrations", "Connect the production/load-code system, reconcile handoffs, define label/registration import, and establish daily exception ownership.", "End-to-end operational handoff is reconciled without manual duplicate entry."),
        ("Phase 4 — Controlled rollout", "Add locations in waves, provision devices, assign managers, monitor freshness/quality, and run weekly pilot reviews.", "Each wave meets its exit metrics before the next wave."),
        ("Phase 5 — Analytics and optimization", "Load history, validate warehouse outlook, add approved holidays/capacity inputs, publish governed exports/Power BI/Fabric products, and tune service recommendations.", "Forecasts and planning reports are trusted enough to guide transportation decisions."),
    ]
    add_table(doc, ["Phase", "Work", "Exit evidence"], phases, [2300, 4800, 2260], font_size=8.8)
    add_heading(doc, "Pilot acceptance measures to define with Goodwill", 2)
    for text in [
        "Scan completion rate by workflow and location, including the percentage of events that finish without a paper fallback.",
        "Offline sync success, maximum acceptable pending age, and time to recover after a connectivity outage.",
        "Rate of review flags, duplicate/sequence issues, and time from review creation to disposition.",
        "Container location accuracy measured by controlled physical counts, not just application state.",
        "Time saved in daily load-code/inventory reconciliation compared with the current paper/spreadsheet process.",
        "Device freshness, app-version compliance, scanner availability, and replacement turnaround.",
        "Manager and corporate user comprehension: can users explain what a status means without technical assistance?",
    ]:
        add_bullet(doc, text, bullet_id)


def add_questions(doc, bullet_id):
    add_heading(doc, "13. Goodwill discovery questions and decisions", 1)
    add_body(doc, "These are the priority questions that materially change data behavior, access control, or pilot success. They should be answered by the operational owner, transportation, store/warehouse leadership, IT/security, and the owner of the production system—not only by a software developer.")
    groups = [
        ("A. Containers, labels, and lifecycle", [
            "Which reusable container types are in scope for the first pilot (bin, cart, gaylord, crate, other), and which goods categories/classifications must be captured?",
            "Are labels unique across the entire Goodwill organization, and who is the authority that creates, prints, replaces, retires, and reissues them?",
            "What is the official lifecycle for a damaged, missing, retired, or newly introduced container?",
            "Do opening counts need to be imported from paper/spreadsheets, or will the pilot establish a clean baseline through scanning?",
        ]),
        ("B. Events and movement", [
            "Exactly when should an employee use Mark full, Send out, Receive, and Mark empty? Are there partial loads, split loads, backhauls, or exceptions?",
            "When a departure has no known destination, what is the expected receiving process and how should transportation follow up on an aging open handoff?",
            "Can the same truck or container make multiple hops, and can a receiving location scan before the departure is visible due to offline delay?",
            "Who is allowed to record a receipt, and what evidence is required when a container arrives without a visible departure record?",
            "What counts as a material correction that needs corporate approval versus a low-risk local correction?",
        ]),
        ("C. Locations and access", [
            "Which sites are stores, Donation Xpress locations, warehouses, or other types, and who owns the authoritative location list?",
            "Should a location manager see only current containers, or also historical departures, open reviews, and local reports?",
            "Can a manager move a scanner between locations, or must corporate control all cross-location moves?",
            "What should happen to devices, containers, managers, load codes, and historical references when a location closes or changes name?",
            "Which corporate roles are Organization Owners, and who is the second owner required for continuity and dual control?",
        ]),
        ("D. Scanner, network, and offline", [
            "Which Android/Unitech model, OS version, scanner interface, camera/QR behavior, charging process, and replacement model will be used?",
            "How often and how long can a device realistically be offline? What is the maximum acceptable local queue size and pending age?",
            "Can shared devices be used by any employee at the assigned site, or must individual employees sign in for every scan?",
            "Will Goodwill use Intune/MDM, managed app distribution, kiosk mode, remote wipe, and certificate/device credential provisioning?",
            "What network controls are required: VPN, private endpoints, allowlisted mobile networks, certificates, or normal internet access?",
        ]),
        ("E. Load codes and integrations", [
            "What production system receives load codes, what fields are required, and which system owns the code format?",
            "Does a load code mean only “marked full,” or is there a separate production acceptance/hand-off event?",
            "How should duplicate, rejected, corrected, or late load-code submissions reconcile with the production system?",
            "Are there existing APIs, files, schedules, or security requirements for the integration?",
        ]),
        ("F. Inventory, service, and forecasts", [
            "What is the official definition of on hand, full, empty, unavailable, in transit, and unknown for each container type?",
            "What minimum/maximum or safety-buffer target does each store/warehouse need for each goods/container type, and who can change it?",
            "What makes a transportation row critical, urgent, or informational? Can a planner override the recommendation and must the override be audited?",
            "Which holidays, events, donation campaigns, weather patterns, or seasonal periods materially change expected volume?",
            "What history is available for trend validation, how is it stored, and which forecast error would be acceptable for a pilot?",
        ]),
        ("G. Security, privacy, and operations", [
            "Which Goodwill Entra tenant, groups, MFA/conditional-access rules, retention policies, and security review process apply?",
            "What data is sensitive, how long must observations/audit records be retained, and who can export them?",
            "What are the required backup/restore objectives, monitoring alerts, incident response path, and support hours?",
            "Will Goodwill require private networking, a specific Azure subscription/resource group, regional restrictions, or a formal change-management process?",
            "Who owns the service after rollout, and what response/maintenance expectations are reasonable given the development/support model?",
        ]),
        ("H. Pilot success", [
            "Which two locations are the first pilot, what workflows must be exercised, and how long will the test run?",
            "What measurable result decides go/no-go: fewer missing containers, scan accuracy, reconciliation time, sync reliability, adoption, or another metric?",
            "Who attends weekly pilot reviews, who can stop the pilot, and how are issues classified and escalated?",
        ]),
    ]
    for heading, questions in groups:
        add_heading(doc, heading, 2)
        for q in questions:
            add_bullet(doc, q, bullet_id)


def add_risks(doc, bullet_id):
    add_heading(doc, "14. Risks, controls, and acceptance criteria", 1)
    add_table(doc, ["Risk", "Why it matters", "Current control / next mitigation"], [
        ("Unknown destination is mistaken for a planned destination", "A guessed route can create false operational confidence.", "Origin-only departure language; destination appears only after receipt; validate multi-hop scenarios with Goodwill."),
        ("Offline devices create delayed or duplicate evidence", "Paper processes often fail during the exact outage the app must withstand.", "Local durable queue, UUID/sequence, idempotency, replay status, pending-age telemetry, and outage testing."),
        ("A correction silently rewrites history", "Trust and auditability are lost if a manager edits the official record directly.", "Append-only ledger, reasoned requests, approval/dual control, audit trail, and later-scan supersession."),
        ("Location/user scope is too broad", "A local manager may see or alter unrelated operations.", "Role/scope model, assigned locations, server-side checks, RLS foundation, and permission tests."),
        ("Forecast is treated as fact", "Synthetic or incomplete history can drive poor dispatch decisions.", "Separate planning inputs from scan history; show calculation; label scenario status; validate against real history."),
        ("Pilot authentication reaches production", "A test password bridge is not Goodwill identity or MFA.", "Explicit test-only boundary; Entra/MFA and secret-management gate before real data."),
        ("Hardware or reference data is not maintained", "A stale assignment or missing label can look like a physical exception.", "Device page, version/freshness signals, location administration, ownership questions, and MDM plan."),
    ], [2200, 3300, 3860], font_size=8.8)
    add_heading(doc, "Definition of ready for a two-location pilot", 2)
    for text in [
        "Goodwill has approved the four core event meanings, the unknown-destination receiving process, and the correction/approval policy.",
        "The chosen scanners can reliably scan labels, report identity/version, work offline, replay, and be remotely managed or replaced.",
        "The API is behind Goodwill-approved identity and network controls; no pilot password or public synthetic route is used for real data.",
        "Opening labels, location references, goods/container types, and user/device assignments are loaded through an agreed process.",
        "Store and warehouse staff can complete the workflow and understand the difference between pending, synced, review, departed, arrived, and empty.",
        "Corporate users can investigate Activity, use Audit, resolve review/correction cases, export the agreed reports, and explain data-health signals.",
        "A rollback/stop-pilot process exists, including how paper records are reconciled if the service is unavailable.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_label_paragraph(doc, "Suggested Goodwill sign-off", "Treat the first pilot as a controlled operational experiment. Approve the rules, measures, and stop conditions before asking employees to rely on the system as the official record.", fill=PALE_TEAL, border=TEAL)


def add_glossary(doc):
    add_heading(doc, "15. Glossary", 1)
    add_table(doc, ["Term", "Plain-language meaning"], [
        ("Observation", "A scanner's record that an action was seen at a time, location, and device."),
        ("Event ledger", "The append-only history of observations and governed administrative decisions."),
        ("Projection", "The current state derived from the evidence, such as full at a location or open in transit."),
        ("Effective time", "The time used for ordering after measured device/server clock offset is considered."),
        ("Received time", "When the API accepted the upload; useful for measuring offline/sync latency."),
        ("Load code", "A validated identifier generated when a container is marked full; its exact production meaning is a Goodwill integration decision."),
        ("Open handoff", "A departure with a confirmed origin and no confirmed receiving scan yet."),
        ("Review case", "A governed work item created when evidence is ambiguous, conflicting, or outside a configured rule."),
        ("Correction", "A new, attributed administrative decision; it does not edit or delete the original observation."),
        ("Activity", "The operational feed of physical scanner observations and movement."),
        ("Audit trail", "The accountability record of administrator sign-ins, controls, approvals, corrections, and reasons."),
        ("Data health", "Evidence-quality signals that help an administrator decide where to investigate."),
        ("Location Manager", "A scoped local operator who can manage assigned-site work but cannot govern the wider network."),
        ("Organization Owner", "A Goodwill corporate administrator with network-wide governance and user/approval responsibility."),
    ], [2200, 7160], font_size=9.4)
    add_heading(doc, "Implementation references", 2)
    add_body(doc, "The working repository contains the deeper engineering references behind this overview: README.md, docs/architecture.md, docs/baseline-assumptions.md, docs/access-control-foundation.md, docs/reporting-foundation.md, docs/location-administration.md, docs/deployment.md, and docs/local-testing.md. Those files are intended for the development/implementation team; this brief is the Goodwill-facing summary and decision aid.")
    add_label_paragraph(doc, "Closing note", "The foundation is intentionally designed to be honest about uncertainty. The next step is not to add more screens for their own sake; it is to validate the operating rules with Goodwill, connect production identity and hardware, and prove the workflow in two real locations with measurable accuracy and recovery outcomes.", fill=PALE_BLUE, border=BLUE)


def main():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        add_header_footer(section)
    bullet_id = define_numbering(doc, "bullet")
    number_id = define_numbering(doc, "number")

    add_cover(doc, bullet_id)
    doc.add_page_break()
    add_contents(doc)
    doc.add_page_break()
    add_exec_summary(doc, bullet_id)
    add_principles(doc, bullet_id)
    add_architecture(doc, bullet_id)
    add_mobile(doc, bullet_id)
    add_admin(doc, bullet_id)
    add_accuracy(doc, bullet_id)
    add_governance(doc, bullet_id, number_id)
    add_reporting(doc, bullet_id)
    add_scenarios(doc, bullet_id)
    add_deployment(doc, bullet_id)
    add_readiness(doc, bullet_id)
    add_roadmap(doc, bullet_id, number_id)
    add_questions(doc, bullet_id)
    add_risks(doc, bullet_id)
    add_glossary(doc)

    core = doc.core_properties
    core.title = "StackTrack Development Overview — Goodwill Reusable Container Operations Pilot"
    core.subject = "Current implementation, accuracy model, governance, roadmap, and Goodwill discovery questions"
    core.author = "Parker Smith / StackTrack development"
    core.keywords = "StackTrack, Goodwill, reusable containers, inventory, audit, offline scanning, pilot"
    core.comments = "Goodwill-facing development overview. Test data only."
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
