from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
p=Path('docs/StackTrack_Goodwill_Development_Overview.docx')
d=Document(p)
text='\n'.join([x.text for x in d.paragraphs]+[c.text for t in d.tables for row in t.rows for c in row.cells])
print('paragraphs',len(d.paragraphs),'tables',len(d.tables),'sections',len(d.sections),'chars',len(text))
for needle in ['â','TODO','TBD','Lorem','placeholder','undefined','needs_review','DeviceSequenceCollision','admin.signed_in']:
 print(needle, text.count(needle))
print('headings', [p.text for p in d.paragraphs if p.style.name.startswith('Heading')][:30])
for i,t in enumerate(d.tables):
 grid=t._tbl.tblGrid
 widths=[int(c.get(qn('w:w'))) for c in grid]
 if sum(widths)!=9360: print('bad table',i,widths,sum(widths))
print('size',p.stat().st_size)
