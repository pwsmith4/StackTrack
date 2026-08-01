from docx import Document
from pathlib import Path
d=Document()
d.add_heading('Hello',1)
d.add_paragraph('A small render test.')
d.save(Path('.tmp/minimal.docx'))
