from zipfile import ZipFile
from xml.etree import ElementTree as ET
from docx import Document
from docx.oxml.ns import qn
p='docs/StackTrack_Goodwill_Development_Overview.docx'
with ZipFile(p) as z:
    for name in ['word/document.xml','word/styles.xml','word/numbering.xml']:
        ET.fromstring(z.read(name)); print('xml ok',name)
d=Document(p)
s=d.sections[0]
print('page inches',s.page_width/914400,s.page_height/914400,'margins',s.left_margin/914400,s.right_margin/914400,s.top_margin/914400,s.bottom_margin/914400)
print('tables',len(d.tables),'paragraphs',len(d.paragraphs))
for i,t in enumerate(d.tables):
    widths=[int(c.get(qn('w:w'))) for c in t._tbl.tblGrid]
    assert sum(widths)==9360,(i,widths,sum(widths))
print('all tables 9360 dxa')
